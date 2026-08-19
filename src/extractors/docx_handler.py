from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
    except Exception as e:
        raise RuntimeError(f"cannot open docx: {e}") from e
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    # also tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append("\t".join(cells))
    text = "\n".join(parts)
    return text


def write_docx(original_path: Path, translated_text: str, output_path: Path) -> Path:
    """
    Best-effort: replace paragraph texts in order with translated lines.
    Layout/styles are preserved as much as python-docx allows.
    If line counts differ, remaining paragraphs are cleared or appended.
    """
    try:
        doc = Document(str(original_path))
    except Exception:
        # create fresh document
        doc = Document()
        for line in translated_text.splitlines():
            doc.add_paragraph(line)
        doc.save(str(output_path))
        return output_path

    lines = translated_text.splitlines()
    paras = list(doc.paragraphs)
    for i, para in enumerate(paras):
        if i < len(lines):
            # clear runs and set new text to keep some style on first run if possible
            if para.runs:
                para.runs[0].text = lines[i]
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.text = lines[i]
        else:
            if para.runs:
                para.runs[0].text = ""
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.text = ""

    # extra lines → new paragraphs
    if len(lines) > len(paras):
        for line in lines[len(paras) :]:
            doc.add_paragraph(line)

    doc.save(str(output_path))
    return output_path
